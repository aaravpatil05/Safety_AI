import docx
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

def replace_text_in_paragraph(paragraph, replacements):
    for key, val in replacements.items():
        if key in paragraph.text:
            # Simple replacement might lose inline formatting, but it's safe for plain placeholders
            inline = paragraph.runs
            for i in range(len(inline)):
                if key in inline[i].text:
                    inline[i].text = inline[i].text.replace(key, val)

def process_document(input_path, output_path):
    doc = docx.Document(input_path)
    
    replacements = {
        "TITLE": "SafetyAI: AI-Powered Personal Safety System",
        "Title": "SafetyAI: AI-Powered Personal Safety System",
        "Guide Name": "Anirudh Chavan",
        "Name of Guide": "Anirudh Chavan",
    }
    
    # Replace in paragraphs
    for para in doc.paragraphs:
        replace_text_in_paragraph(para, replacements)
        
    # Replace in tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    replace_text_in_paragraph(para, replacements)
                    
    # The template has "Student Name" and "PRN Number" placeholders.
    # We will search for them and replace them.
    students = [
        "Aarav Patil - 2023000057",
        "Ayush Mane - 2023000135",
        "Tanmay Joshi - 2023000137",
        "Sneha Gaikwad - 2023000075"
    ]
    student_text = "\n".join(students)
    
    for para in doc.paragraphs:
        if "Student Name" in para.text:
            for r in para.runs:
                r.text = r.text.replace("Student Name", students[0])
            # Add other students
            para.insert_paragraph_before(students[1])
            para.insert_paragraph_before(students[2])
            para.insert_paragraph_before(students[3])
        if "PRN Number" in para.text:
            for r in para.runs:
                r.text = r.text.replace("PRN Number", "")

    # Now let's fill in the chapters
    # We will look for headings like "Chapter 1: Abstract" and insert content after them.
    
    content = {
        "Chapter 1: Abstract": [
            "Problem Statement: In modern cities, ensuring personal safety during emergencies is a critical challenge. Existing emergency systems often require manual intervention, which is not always possible during a crisis.",
            "Objective: To develop an intelligent, automated personal safety application (SafetyAI) that provides real-time tracking, AI-driven voice SOS, geofencing, and automated alert systems without requiring active user interaction in danger zones.",
            "Technology Used: Android (Java), Google Fused Location API, osmdroid (Maps), Android Speech Recognizer, Local/Cloud Servers.",
            "Methodology: The system uses background services to continuously monitor voice commands (shouting 'Help') and location. Upon triggering an SOS, it silently records audio evidence, uploads it to a secure server, and dispatches SMS alerts with live tracking links to predefined contacts.",
            "Expected Outcome: A robust, network-resilient application capable of notifying authorities and emergency contacts instantly, accompanied by undeniable audio evidence and exact GPS coordinates."
        ],
        "Chapter 2: Introduction": [
            "Background of the Project:",
            "Personal security is a growing concern globally. While smartphones are ubiquitous, current safety applications require multiple taps to trigger an alert, which is impractical in sudden, threatening situations.",
            "Motivation:",
            "The motivation behind SafetyAI stems from the need for a hands-free, autonomous safety net. By leveraging AI voice recognition and smart geofencing, the phone itself becomes an active guardian rather than just a passive communication tool.",
            "Need of the System:",
            "Victims of emergencies often cannot reach their phones. A system that detects distress through voice or enters high-risk zones automatically can save lives by drastically reducing emergency response times.",
            "Problem Definition:",
            "Design and implement a mobile application that can autonomously detect emergencies, capture forensic evidence (audio), and broadcast a live dashboard containing location and audio data to emergency contacts, even under poor network conditions.",
            "Project Scope:",
            "The project encompasses an Android application with background tracking, offline SMS fallback, live server-side dashboards for evidence viewing, and dynamic risk maps for cities."
        ],
        "Chapter 3: Requirement Analysis": [
            "3.1 Requirement Analysis",
            "Functional Requirements:",
            "- Voice SOS: System must trigger SOS upon hearing 'Help' multiple times.",
            "- Location Tracking: System must track precise GPS coordinates in the background.",
            "- SMS Broadcasting: System must send SMS to 3 contacts without internet.",
            "- Evidence Upload: System must record microphone audio and upload it to a server.",
            "- Risk Map: System must display city-level risk zones (e.g., Mahim, CSMT).",
            "Non-Functional Requirements:",
            "- Reliability: Must function reliably under lock-screen and background states.",
            "- Performance: Location must be fetched within 3 seconds of an emergency.",
            "- Security: Audio evidence must be securely handled.",
            "3.2 Feasibility Study",
            "Technical Feasibility: High. Android provides the necessary hardware access (GPS, Mic) and APIs.",
            "Economic Feasibility: High. The system utilizes existing smartphone hardware without requiring extra devices.",
            "Operational Feasibility: High. The app integrates seamlessly into daily user routines via background monitoring."
        ],
        "Chapter 4: System Design": [
            "4.1 Architecture Diagram",
            "The system follows a Client-Server architecture. The Android client handles sensor inputs (Location, Voice) and uploads payloads to the Python backend server, which hosts a live dashboard for emergency contacts.",
            "4.2 UML Diagrams",
            "Use Case Diagram: User configures contacts -> System monitors environment -> System detects threat -> System alerts contacts.",
            "Sequence Diagram: Voice Input -> Trigger SOS -> Get Location -> Record Audio -> Upload to Server -> Send SMS -> Contacts view Dashboard.",
            "Activity Diagram: App start -> Load preferences -> Start background listeners -> If threat detected -> Execute SOS Protocol.",
            "4.3 Database Design",
            "Data is stored locally using SharedPreferences for user settings (contacts, sensitivity).",
            "Server-side uses JSON-based flat files for mapping Evidence IDs to timestamps and GPS coordinates."
        ],
        "Chapter 5: Implementation": [
            "Module Description:",
            "1. Location Module: Uses FusedLocationProviderClient to enforce fresh GPS fixes, bypassing cached data for strict accuracy.",
            "2. SOS Trigger Module: Listens via SpeechRecognizer. If 'Help' is detected 7 times, it triggers the SOS flag.",
            "3. Evidence Module: Uses MediaRecorder to capture audio in 3GPP format, securing it in local storage before attempting network upload.",
            "4. Networking Module: Features auto-discovery (UDP) and SSH tunneling to expose local servers to the public internet.",
            "Algorithm:",
            "1. Initialize Speech Recognizer.",
            "2. Loop: Listen to audio.",
            "3. If text contains 'help', increment counter.",
            "4. If counter >= 7, invoke toggleSos()."
        ],
        "Chapter 6 : Testing": [
            "7.1 Test Plan",
            "The system is tested under various constraints: offline mode, background mode, and different locations.",
            "7.2 Test Cases Table",
            "- TC01: Voice Trigger | Input: Say 'Help' 7 times | Expected: SOS Activates | Result: Pass",
            "- TC02: Location Fetch | Input: Trigger SOS | Expected: Fresh Location Acquired | Result: Pass",
            "- TC03: SMS Fallback | Input: Disable WiFi/Data, Trigger SOS | Expected: Offline SMS Sent | Result: Pass",
            "- TC04: Risk Map | Input: Open Map | Expected: Mumbai/Maharashtra specific zones shown | Result: Pass",
            "7.3 Types of Testing",
            "- Unit Testing: Validating individual modules like SMS formatting and location filtering.",
            "- Integration Testing: Ensuring audio upload maps correctly to the live dashboard.",
            "- System Testing: End-to-end testing from Voice SOS to Dashboard viewing.",
            "- User Acceptance Testing: Field testing in different network conditions."
        ],
        "Chapter 7: Future Scope": [
            "Future enhancements for SafetyAI include:",
            "- Direct integration with police dispatch systems (Dial 112).",
            "- Live video streaming in addition to audio recording.",
            "- Fall detection utilizing accelerometer data to detect sudden impacts.",
            "- Cross-platform support for iOS devices."
        ],
        "Chapter 8: Conclusion": [
            "Summary of Work:",
            "SafetyAI successfully transforms a standard smartphone into an autonomous emergency guardian.",
            "Objectives Achieved:",
            "We achieved hands-free voice SOS, accurate live location tracking, secure evidence handling, and robust offline SMS alerts.",
            "Key Learnings:",
            "We learned the importance of bypassing cached location data on aggressive Android OEM skins (like OnePlus) and the complexities of maintaining stable background services."
        ],
        "Chapter 9: References": [
            "1. Android Developers Documentation: FusedLocationProviderClient.",
            "2. osmdroid Documentation: Implementing offline/online maps.",
            "3. Python HTTP Server Documentation.",
            "4. Localhost.run SSH Tunneling Guide."
        ]
    }
    
    # Simple insertion logic:
    # Find paragraph containing the chapter title, then insert the content paragraphs below it.
    
    for i, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        for chapter, lines in content.items():
            if chapter in text:
                # Add page break before chapter heading to ensure it's on a fresh page
                # If it's the first chapter, it might already be on a fresh page, but we'll enforce it.
                if "Chapter 1: Abstract" in text:
                     # It's usually good to keep the existing page structure if it's already there
                     pass
                else:
                     para.insert_paragraph_before("").add_run().add_break(docx.enum.text.WD_BREAK.PAGE)
                
                # Make heading smaller (the user requested headings should be small, not so big)
                for run in para.runs:
                    run.font.size = Pt(14)
                    run.font.bold = True
                
                # Insert content lines
                for line in lines:
                    new_p = para.insert_paragraph_before(line)
                    new_p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                    for r in new_p.runs:
                        r.font.size = Pt(12)
    
    # Save the formatted document
    output_filepath = "/Users/aaravpatil/Documents/SafetyAI_Capstone_Report.docx"
    doc.save(output_filepath)
    print(f"Document saved to {output_filepath}")

if __name__ == '__main__':
    process_document('/Users/aaravpatil/Documents/capstone report format.docx', '')
