import docx
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
import copy

def delete_paragraph(paragraph):
    p = paragraph._element
    p.getparent().remove(p)
    p._p = p._element = None

def process_document(input_path, output_path):
    doc = docx.Document(input_path)
    
    # 1. Replace plain text placeholders
    replacements = {
        "TITLE": "SafetyAI - AI-Powered Personal Safety & Emergency Response System",
        "Title": "SafetyAI - AI-Powered Personal Safety & Emergency Response System",
        "Guide Name": "Anirudh Chavan",
        "Name of Guide": "Anirudh Chavan",
    }
    for para in doc.paragraphs:
        for key, val in replacements.items():
            if key in para.text:
                for r in para.runs:
                    if key in r.text:
                        r.text = r.text.replace(key, val)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    for key, val in replacements.items():
                        if key in para.text:
                            for r in para.runs:
                                if key in r.text:
                                    r.text = r.text.replace(key, val)

    # 2. Fill Students Table
    # Table 0 is assumed to be Student Name and PRN
    students = [
        ("Aarav Patil", "2023000057"),
        ("Ayush Mane", "2023000135"),
        ("Tanmay Joshi", "2023000137"),
        ("Sneha Gaikwad", "2023000075")
    ]
    t = doc.tables[0]
    # Fill existing rows (1, 2, 3)
    t.cell(1, 0).text = students[0][0]
    t.cell(1, 1).text = students[0][1]
    t.cell(2, 0).text = students[1][0]
    t.cell(2, 1).text = students[1][1]
    t.cell(3, 0).text = students[2][0]
    t.cell(3, 1).text = students[2][1]
    # Add a row for the 4th student
    row = t.add_row()
    row.cells[0].text = students[3][0]
    row.cells[1].text = students[3][1]

    # 3. Replace pointers with actual content
    content_map = {
        "Problem Statement": "Problem Statement\nIn modern cities, ensuring personal safety during emergencies is a critical challenge. Existing emergency systems often require manual intervention, which is not always possible during a crisis.",
        "Objective": "Objective\nTo develop an intelligent, automated personal safety application (SafetyAI) that provides real-time tracking, AI-driven voice SOS, geofencing, and automated alert systems without requiring active user interaction in danger zones.",
        "Technology Used": "Technology Used\nAndroid (Java), Google Fused Location API, osmdroid (Maps), Android Speech Recognizer, Local/Cloud Python Servers.",
        "Methodology": "Methodology\nThe system uses background services to continuously monitor voice commands (shouting 'Help') and location. Upon triggering an SOS, it silently records audio evidence, uploads it to a secure server, and dispatches SMS alerts with live tracking links to predefined contacts.",
        "Expected Outcome": "Expected Outcome\nA robust, network-resilient application capable of notifying authorities and emergency contacts instantly, accompanied by undeniable audio evidence and exact GPS coordinates.",
        
        "Background of the Project": "Background of the Project\nPersonal security is a growing concern globally. While smartphones are ubiquitous, current safety applications require multiple taps to trigger an alert, which is impractical in sudden, threatening situations.",
        "Motivation": "Motivation\nThe motivation behind SafetyAI stems from the need for a hands-free, autonomous safety net. By leveraging AI voice recognition and smart geofencing, the phone itself becomes an active guardian rather than just a passive communication tool.",
        "Need of the System": "Need of the System\nVictims of emergencies often cannot reach their phones. A system that detects distress through voice or enters high-risk zones automatically can save lives by drastically reducing emergency response times.",
        "Problem Definition": "Problem Definition\nDesign and implement a mobile application that can autonomously detect emergencies, capture forensic evidence (audio), and broadcast a live dashboard containing location and audio data to emergency contacts, even under poor network conditions.",
        "Project Scope": "Project Scope\nThe project encompasses an Android application with background tracking, offline SMS fallback, live server-side dashboards for evidence viewing, and dynamic risk maps for cities.",
        
        "Functional Requirements": "Functional Requirements\n- Voice SOS: System must trigger SOS upon hearing 'Help' multiple times.\n- Location Tracking: System must track precise GPS coordinates in the background.\n- SMS Broadcasting: System must send SMS to 3 contacts without internet.\n- Evidence Upload: System must record microphone audio and upload it to a server.",
        "Non-Functional Requirements": "Non-Functional Requirements\n- Reliability: Must function reliably under lock-screen and background states.\n- Performance: Location must be fetched within 3 seconds of an emergency.\n- Security: Audio evidence must be securely handled.",
        "Technical Feasibility": "Technical Feasibility\nHigh. Android provides the necessary hardware access (GPS, Mic) and APIs for background processing.",
        "Economic Feasibility": "Economic Feasibility\nHigh. The system utilizes existing smartphone hardware without requiring extra specialized devices.",
        "Operational Feasibility": "Operational Feasibility\nHigh. The app integrates seamlessly into daily user routines via autonomous background monitoring.",
        
        "4.1 Architecture Diagram": "4.1 Architecture Diagram\nThe system follows a Client-Server architecture. The Android client handles sensor inputs (Location, Voice) and uploads payloads to the Python backend server, which hosts a live dashboard for emergency contacts.",
        "(Insert Block Diagram)": "The block diagram illustrates the flow from the Android Client's Sensors (Mic, GPS) to the Local/Cloud Server for evidence hosting, and finally to Emergency Contacts via SMS.",
        "4.2 UML Diagrams": "4.2 UML Diagrams\nDetailed UML models are used to map out the interactions between the user, the AI background services, and the external APIs.",
        "Use Case Diagram": "Use Case Diagram\nUser configures contacts -> System monitors environment -> System detects threat -> System alerts contacts.",
        "Class Diagram": "Class Diagram\nMainActivity manages services -> SpeechRecognizer handles audio -> FusedLocationClient handles GPS.",
        "Sequence Diagram": "Sequence Diagram\nVoice Input -> Trigger SOS -> Get Location -> Record Audio -> Upload to Server -> Send SMS -> Contacts view Dashboard.",
        "Activity Diagram": "Activity Diagram\nApp start -> Load preferences -> Start background listeners -> If threat detected -> Execute SOS Protocol.",
        "4.3 Database Design": "4.3 Database Design\nData is stored locally using SharedPreferences for user settings (contacts, sensitivity).",
        "ER Diagram": "ER Diagram\nEntities include User, EmergencyContact, and EvidenceRecord.",
        "Table Structure": "Table Structure\nEvidence table: ID (Timestamp), Latitude, Longitude, Audio Filename.",
        "Data Dictionary": "Data Dictionary\nDescribes the structure of the JSON payload sent to the server.",
        
        "Module Description": "Module Description\n1. Location Module: Uses FusedLocationProviderClient to enforce fresh GPS fixes.\n2. SOS Trigger Module: Listens via SpeechRecognizer. If 'Help' is detected 7 times, it triggers the SOS flag.\n3. Evidence Module: Uses MediaRecorder to capture audio in 3GPP format.\n4. Networking Module: Features auto-discovery (UDP) and SSH tunneling.",
        "Algorithm (if any)": "Algorithm\n1. Initialize Speech Recognizer.\n2. Loop: Listen to audio.\n3. If text contains 'help', increment counter.\n4. If counter >= 7, invoke toggleSos().",
        "Flowcharts": "Flowcharts\nVisual representation of the background service lifecycle.",
        "Screenshots of System": "Screenshots of System\n(Screenshots of the Home Screen, Risk Map, and Emergency Dashboard will be attached here)",
        "Code Snippets (Optional in Appendix)": "Code Snippets\nKey algorithms for UDP discovery and location fetching are included in the appendix.",
        
        "7.1 Test Plan": "7.1 Test Plan\nThe system is tested under various constraints: offline mode, background mode, and different locations.",
        "7.3 Types of Testing": "7.3 Types of Testing\nWe conducted rigorous testing across the software lifecycle.",
        "Unit Testing": "Unit Testing\nValidating individual modules like SMS formatting and location filtering.",
        "Integration Testing": "Integration Testing\nEnsuring audio upload maps correctly to the live dashboard on the Python server.",
        "System Testing": "System Testing\nEnd-to-end testing from Voice SOS to Dashboard viewing on a target device (OnePlus Nord 2).",
        "User Acceptance Testing": "User Acceptance Testing\nField testing in different network conditions and simulated emergencies.",
        
        "Summary of Work": "Summary of Work\nSafetyAI successfully transforms a standard smartphone into an autonomous emergency guardian.",
        "Objectives Achieved": "Objectives Achieved\nWe achieved hands-free voice SOS, accurate live location tracking, secure evidence handling, and robust offline SMS alerts.",
        "Key Learnings": "Key Learnings\nWe learned the importance of bypassing cached location data on aggressive Android OEM skins and the complexities of maintaining stable background services."
    }

    # Iterate backwards so we can delete safely
    for para in list(doc.paragraphs):
        text = para.text.strip()
        
        # 1. Delete empty paragraphs to fix the "blank page" issue
        if not text:
            delete_paragraph(para)
            continue
            
        # 2. Add page breaks before Chapters
        if text.startswith("Chapter"):
            # Set font size smaller
            for r in para.runs:
                r.font.size = Pt(14)
            # Insert page break before the chapter (unless it's chapter 1 and already at top)
            if "Chapter 1" not in text:
                # Insert a new paragraph before this one with a page break
                new_p = para.insert_paragraph_before("")
                new_p.add_run().add_break(WD_BREAK.PAGE)
        
        # 3. Replace pointers with actual text
        for key, val in content_map.items():
            if text == key or text == f"1. {key}" or text == f"2. {key}" or text == f"3. {key}" or text == f"4. {key}" or text == f"5. {key}":
                para.text = val
                para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                for r in para.runs:
                    r.font.size = Pt(12)
                break

    # Save
    doc.save(output_path)

if __name__ == '__main__':
    process_document('/Users/aaravpatil/Documents/capstone report format.docx', '/Users/aaravpatil/Documents/SafetyAI_Capstone_Report_V2.docx')
